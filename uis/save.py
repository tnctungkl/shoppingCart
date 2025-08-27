from pathlib import Path
import json
from openpyxl import Workbook
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document

def _ensure_dir(path: str):
    Path(Path(path).parent).mkdir(parents=True, exist_ok=True)

def save_json(cart_snapshot: dict, filename: str = "saved/cart.json"):
    _ensure_dir(filename)
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(cart_snapshot, f, indent=4, ensure_ascii=False)
    return filename

def save_excel(cart_snapshot: dict, filename: str = "saved/cart.xlsx"):
    _ensure_dir(filename)
    wb = Workbook()
    ws = wb.active
    ws.title = "Cart"
    ws.append(["Product ID", "Name", "Quantity", "Price", "Shipping", "Subtotal"])
    for it in cart_snapshot.get("items", []):
        ws.append([
            it["product_id"], it["name"], it["quantity"], it["price"], it["shipping"], it["subtotal"]
        ])
    ws.append([])
    ws.append(["", "", "", "", "TOTAL", cart_snapshot.get("total", 0)])
    wb.save(filename)
    return filename

def save_pdf(cart_snapshot: dict, filename: str = "saved/cart.pdf"):
    _ensure_dir(filename)
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    canSave = []
    canSave.append(Paragraph("<b>Tungshoop – Cart Save Section</b>", styles["Title"]))
    canSave.append(Spacer(1, 12))

    data = [["Product ID", "Name", "Quantity", "Price", "Shipping", "Subtotal"]]
    for it in cart_snapshot.get("items", []):
        data.append([
            it["product_id"], it["name"], it["quantity"], f"₺{it['price']}", f"₺{it['shipping']}", f"₺{it['subtotal']}"
        ])
    data.append(["", "", "", "", "TOTAL", f"₺{cart_snapshot.get('total', 0)}"])

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('ALIGN', (-2,1), (-1,-1), 'RIGHT'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
    ]))

    canSave.append(table)
    doc.build(canSave)
    return filename

def save_docx(cart_snapshot: dict, filename: str = "saved/cart.docx"):
    _ensure_dir(filename)
    doc = Document()
    doc.add_heading("Tungshoop – Save Section", 0)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Product ID"; hdr[1].text = "Name"; hdr[2].text = "Quantity"; hdr[3].text = "Price"; hdr[4].text = "Shipping"; hdr[5].text = "Subtotal"
    for it in cart_snapshot.get("items", []):
        row = table.add_row().cells
        row[0].text = str(it["product_id"]) ; row[1].text = str(it["name"]) ; row[2].text = str(it["quantity"]) ; row[3].text = f"₺{it['price']}" ; row[4].text = f"₺{it['shipping']}" ; row[5].text = f"₺{it['subtotal']}"
    row = table.add_row().cells
    row[0].text = ""; row[1].text = ""; row[2].text = ""; row[3].text = ""; row[4].text = "TOTAL"; row[5].text = f"₺{cart_snapshot.get('total', 0)}"
    doc.save(filename)
    return filename